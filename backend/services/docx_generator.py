"""DOCX workbook generator — produces professional math workbooks with bordered question boxes,
OMML math rendering, and precise page budget management.

Reference format: thick-bordered question boxes with beige backgrounds, option letter boxes,
3-column page headers, colored section titles, and proper Office Math (OMML) for equations.
"""

import logging
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Emu, Inches
from lxml import etree

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
# OMML Math Rendering Engine
# ═══════════════════════════════════════════════════════════════════════════════

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _omml(tag: str) -> OxmlElement:
    """Create an OMML element with the math namespace."""
    el = OxmlElement(f"m:{tag}")
    return el


def _omml_run(text: str, italic: bool = True, bold: bool = False) -> OxmlElement:
    """Create an OMML math run (m:r) containing text."""
    r = _omml("r")
    # Run properties
    rPr = _omml("rPr")
    sty = _omml("sty")
    if italic and not bold:
        sty.set(qn("m:val"), "i")
    elif bold and not italic:
        sty.set(qn("m:val"), "b")
    elif bold and italic:
        sty.set(qn("m:val"), "bi")
    else:
        sty.set(qn("m:val"), "p")
    rPr.append(sty)
    r.append(rPr)
    # Word run properties for font
    wRPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Cambria Math")
    rFonts.set(qn("w:hAnsi"), "Cambria Math")
    wRPr.append(rFonts)
    r.append(wRPr)
    # Text element
    t = OxmlElement("m:t")
    t.text = text
    # Preserve spaces
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _omml_fraction(numerator_elements: list, denominator_elements: list) -> OxmlElement:
    """Create an OMML fraction (m:f) element."""
    f = _omml("f")
    # Fraction properties
    fPr = _omml("fPr")
    f.append(fPr)
    # Numerator
    num = _omml("num")
    for el in numerator_elements:
        num.append(el)
    f.append(num)
    # Denominator
    den = _omml("den")
    for el in denominator_elements:
        den.append(el)
    f.append(den)
    return f


def _omml_radical(content_elements: list, degree_elements: list = None) -> OxmlElement:
    """Create an OMML radical/square root (m:rad) element."""
    rad = _omml("rad")
    # Radical properties
    radPr = _omml("radPr")
    if not degree_elements:
        # Hide degree for square root
        degHide = _omml("degHide")
        degHide.set(qn("m:val"), "1")
        radPr.append(degHide)
    rad.append(radPr)
    # Degree
    deg = _omml("deg")
    if degree_elements:
        for el in degree_elements:
            deg.append(el)
    rad.append(deg)
    # Content under radical
    e = _omml("e")
    for el in content_elements:
        e.append(el)
    rad.append(e)
    return rad


def _omml_superscript(base_elements: list, sup_elements: list) -> OxmlElement:
    """Create an OMML superscript (m:sSup) element."""
    sSup = _omml("sSup")
    sSupPr = _omml("sSupPr")
    sSup.append(sSupPr)
    # Base
    e = _omml("e")
    for el in base_elements:
        e.append(el)
    sSup.append(e)
    # Superscript
    sup = _omml("sup")
    for el in sup_elements:
        sup.append(el)
    sSup.append(sup)
    return sSup


def _omml_subscript(base_elements: list, sub_elements: list) -> OxmlElement:
    """Create an OMML subscript (m:sSub) element."""
    sSub = _omml("sSub")
    sSubPr = _omml("sSubPr")
    sSub.append(sSubPr)
    # Base
    e = _omml("e")
    for el in base_elements:
        e.append(el)
    sSub.append(e)
    # Subscript
    sub = _omml("sub")
    for el in sub_elements:
        sub.append(el)
    sSub.append(sub)
    return sSub


def _omml_parentheses(content_elements: list, open_char: str = "(", close_char: str = ")") -> OxmlElement:
    """Create an OMML delimited expression (m:d) with parentheses."""
    d = _omml("d")
    dPr = _omml("dPr")
    begChr = _omml("begChr")
    begChr.set(qn("m:val"), open_char)
    dPr.append(begChr)
    endChr = _omml("endChr")
    endChr.set(qn("m:val"), close_char)
    dPr.append(endChr)
    d.append(dPr)
    e = _omml("e")
    for el in content_elements:
        e.append(el)
    d.append(e)
    return d


class MathParser:
    """Parses simple math notation and produces OMML elements.

    Supported patterns:
    - x^2, x^{10} → superscript
    - x_1, x_{12} → subscript
    - sqrt(expr) → radical
    - frac(a,b) → fraction
    - Unicode symbols pass through: ≠, ≥, ≤, ∈, ∴, ≈, ∠, ±, ×, ÷
    """

    # Regex patterns for math detection
    SUPERSCRIPT_RE = re.compile(r"(\w+)\^(\{[^}]+\}|\d+)")
    SUBSCRIPT_RE = re.compile(r"(\w+)_(\{[^}]+\}|\d+)")
    FRACTION_RE = re.compile(r"frac\(([^,]+),([^)]+)\)")
    SQRT_RE = re.compile(r"sqrt\(([^)]+)\)")
    SIMPLE_FRAC_RE = re.compile(r"(\d+)/(\d+)")

    @classmethod
    def has_math(cls, text: str) -> bool:
        """Check if text contains math patterns that need OMML rendering."""
        if not text:
            return False
        return bool(
            cls.SUPERSCRIPT_RE.search(text)
            or cls.SUBSCRIPT_RE.search(text)
            or cls.FRACTION_RE.search(text)
            or cls.SQRT_RE.search(text)
        )

    @classmethod
    def parse_to_omml(cls, text: str) -> OxmlElement:
        """Parse a math expression string and return an m:oMath element."""
        oMath = _omml("oMath")
        cls._parse_expression(text, oMath)
        return oMath

    @classmethod
    def _parse_expression(cls, text: str, parent: OxmlElement) -> None:
        """Recursively parse expression and append OMML elements to parent."""
        if not text:
            return

        # Try fraction pattern: frac(num, den)
        frac_match = cls.FRACTION_RE.search(text)
        if frac_match:
            before = text[:frac_match.start()]
            after = text[frac_match.end():]
            if before.strip():
                parent.append(_omml_run(before))
            num_text = frac_match.group(1).strip()
            den_text = frac_match.group(2).strip()
            num_elements = [_omml_run(num_text, italic=True)]
            den_elements = [_omml_run(den_text, italic=True)]
            parent.append(_omml_fraction(num_elements, den_elements))
            if after.strip():
                cls._parse_expression(after, parent)
            return

        # Try sqrt pattern: sqrt(content)
        sqrt_match = cls.SQRT_RE.search(text)
        if sqrt_match:
            before = text[:sqrt_match.start()]
            after = text[sqrt_match.end():]
            if before.strip():
                parent.append(_omml_run(before))
            content = sqrt_match.group(1).strip()
            content_elements = [_omml_run(content, italic=True)]
            parent.append(_omml_radical(content_elements))
            if after.strip():
                cls._parse_expression(after, parent)
            return

        # Try superscript pattern: base^exp or base^{exp}
        sup_match = cls.SUPERSCRIPT_RE.search(text)
        if sup_match:
            before = text[:sup_match.start()]
            after = text[sup_match.end():]
            if before.strip():
                parent.append(_omml_run(before))
            base = sup_match.group(1)
            exp = sup_match.group(2).strip("{}")
            base_elements = [_omml_run(base, italic=True)]
            sup_elements = [_omml_run(exp, italic=True)]
            parent.append(_omml_superscript(base_elements, sup_elements))
            if after.strip():
                cls._parse_expression(after, parent)
            return

        # Try subscript pattern: base_sub or base_{sub}
        sub_match = cls.SUBSCRIPT_RE.search(text)
        if sub_match:
            before = text[:sub_match.start()]
            after = text[sub_match.end():]
            if before.strip():
                parent.append(_omml_run(before))
            base = sub_match.group(1)
            sub = sub_match.group(2).strip("{}")
            base_elements = [_omml_run(base, italic=True)]
            sub_elements = [_omml_run(sub, italic=True)]
            parent.append(_omml_subscript(base_elements, sub_elements))
            if after.strip():
                cls._parse_expression(after, parent)
            return

        # No special patterns — plain math text
        parent.append(_omml_run(text, italic=True))


# ═══════════════════════════════════════════════════════════════════════════════
# Color Scheme Constants
# ═══════════════════════════════════════════════════════════════════════════════

class Colors:
    """Color constants matching the reference design."""
    HEADER_GREEN = "C6E0B4"           # Light green for header bar
    SECTION_TITLE_BG = "C6E0B4"       # Light green/yellow highlight
    SECTION_TITLE_TEXT = RGBColor(0x80, 0x00, 0x80)  # Dark purple/maroon
    QUESTION_BOX_BORDER = "000000"    # Black thick border
    QUESTION_BOX_BG = "FFF8EB"        # Light beige/cream
    OPTION_LETTER_BORDER = "000000"   # Black border for letter box
    OPTION_LETTER_BG = "FFFFFF"       # White background
    TIP_BOX_BG = "DCEBFA"            # Light blue
    SEPARATOR_GREEN = "70AD47"        # Green separator line
    COVER_TITLE = RGBColor(0x1A, 0x1A, 0x1A)
    COVER_SUBTITLE = RGBColor(0x55, 0x55, 0x55)

    # Study Book / Solved Examples colors
    EXAMPLE_TITLE_TEXT = RGBColor(0x00, 0x6B, 0x3F)   # Dark green for example titles
    EXAMPLE_BOX_BORDER = "2F5F8F"     # Dark blue border for solution box
    EXAMPLE_BOX_BG = "FFFFFF"         # White background for solution box
    FORMULA_BOX_BG = "FFC8C8"         # Pink/salmon for key formula box
    FORMULA_BOX_BORDER = "CC4444"     # Red border for formula box
    COEFFICIENTS_BOX_BG = "DCEBFA"    # Light blue for coefficients box
    COEFFICIENTS_BOX_BORDER = "4472C4"  # Blue border for coefficients box
    EXAMPLE_NUMBER_BG = "2F5F8F"      # Dark blue background for example number
    EXAMPLE_SECTION_BG = "E8F5E9"     # Very light green for section header

    # Lesson Illustration colors
    LESSON_TITLE_BG = "C8E6C9"        # Light green background bar for lesson title
    LESSON_TITLE_TEXT = RGBColor(0x1B, 0x5E, 0x20)  # Dark green text for lesson title
    LESSON_CONCEPTS_BG = "E8F5E9"     # Very light green for key concepts box
    LESSON_CONCEPTS_BORDER = "4CAF50"  # Green border for concepts
    LESSON_CONCEPTS_HEADER = RGBColor(0x2E, 0x7D, 0x32)  # Green header text
    LESSON_THEOREMS_BORDER = "757575"  # Gray separator for theorems
    LESSON_FORMULAS_BG = "E3F2FD"     # Light blue for formulas box
    LESSON_FORMULAS_BORDER = "1976D2"  # Blue border for formulas
    LESSON_FORMULAS_HEADER = RGBColor(0x0D, 0x47, 0xA1)  # Dark blue header text
    LESSON_NOTES_BG = "FFF8E1"        # Light yellow for important notes
    LESSON_NOTES_BORDER = "F57C00"    # Orange border for notes
    LESSON_NOTES_HEADER = RGBColor(0xE6, 0x51, 0x00)  # Orange header text


# ═══════════════════════════════════════════════════════════════════════════════
# Main Generator Class
# ═══════════════════════════════════════════════════════════════════════════════

class DocxGenerator:
    """Generates a professional .docx math workbook with bordered question boxes,
    OMML math rendering, and page budget management.

    Produces output matching the reference format:
    - 3-column page header with green separator
    - Highlighted section titles in purple
    - Thick-bordered question boxes with beige background
    - Option letter boxes for MCQ
    - Tip sections in light blue boxes
    - Proper OMML math for equations
    """

    # ─── Page Budget Constants ────────────────────────────────────────────────
    _USABLE_HEIGHT_CM = 24.7  # A4 with 2.5cm margins top/bottom

    EXERCISE_HEIGHT_CM = {
        "multiple_choice": 5.5,
        "fill_in_blank": 4.0,
        "fill_blank": 4.0,
        "long_answer": 8.0,
        "show_work": 7.0,
        "word_problem": 6.5,
        "word_problems": 6.5,
        "true_false": 3.5,
        "matching": 7.0,
    }

    EXERCISES_PER_PAGE = {
        "multiple_choice": 3,
        "fill_in_blank": 4,
        "fill_blank": 4,
        "long_answer": 2,
        "show_work": 2,
        "word_problem": 2,
        "word_problems": 2,
        "true_false": 5,
        "matching": 2,
    }

    _ENGLISH_OPTIONS = ["A", "B", "C", "D", "E", "F"]

    def __init__(
        self,
        config: dict,
        exercises: list[dict],
        output_path: str,
        illustration_content: list[dict] | None = None,
        num_pages: int | None = None,
        lesson_illustrations: list[dict] | None = None,
    ) -> None:
        """
        Args:
            config: The full WorkbookConfig as a dict.
            exercises: List of generated exercise objects from LLM.
            output_path: Where to save the .docx file.
            illustration_content: Optional list of LLM-generated solved example dicts.
                When provided (study book mode), these are rendered before exercises.
            num_pages: Target page count (overrides config if provided).
            lesson_illustrations: Optional list of lesson illustration dicts.
                When provided (study book mode), these are rendered before solved examples.
        """
        self.config = config
        self.exercises = exercises or []
        self.output_path = Path(output_path)
        # Solved examples for study book mode (LLM-generated, not raw OCR)
        self.solved_examples = illustration_content or []
        # Lesson illustrations for study book mode (textbook-style summaries)
        self.lesson_illustrations = lesson_illustrations or []
        self.doc = Document()

        # Config shortcuts
        self.formatting = config.get("formatting", {})
        self.structure = config.get("structure", {})
        self.scope = config.get("scope", {})
        self.exercise_config = config.get("exercises", {})

        # Language — default English for this professional format
        self.language = self.formatting.get("language", "english")

        # Font settings
        self._font_name = "Arial"
        self._font_size = 11

        # Page budget
        self._target_pages = num_pages or self.structure.get("total_pages", 20)
        self._include_cover = self.structure.get("include_cover", True)
        self._content_pages = self._target_pages - (1 if self._include_cover else 0)
        self._spacing_scale = 1.0

        # Title and metadata
        self._title = self.formatting.get("title", "Math Workbook")
        self._school_name = self.formatting.get("school_name", "Mathematics Department")
        self._academic_year = self.formatting.get("academic_year", "2024-2025")
        self._grade = self.formatting.get("grade", "")
        self._term = self.formatting.get("term", "")

    # ═══════════════════════════════════════════════════════════════════════════
    # Main Generation Entry Point
    # ═══════════════════════════════════════════════════════════════════════════

    def generate(self) -> str:
        """Main generation method. Returns the output file path."""
        self._setup_page()
        self._set_document_defaults()
        self._calculate_spacing_scale()

        if self._include_cover:
            self._add_cover_page()

        self._add_workbook_content()

        # Add running headers to all sections
        self._add_page_headers()
        self._add_page_numbers()

        # Ensure output directory exists and save
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(self.output_path))

        logger.info(f"Workbook saved to: {self.output_path}")
        return str(self.output_path)

    # ═══════════════════════════════════════════════════════════════════════════
    # Page Setup & Document Defaults
    # ═══════════════════════════════════════════════════════════════════════════

    def _setup_page(self) -> None:
        """Configure A4 page size and margins."""
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    def _set_document_defaults(self) -> None:
        """Set document-level default font and paragraph styles."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self._font_name
        font.size = Pt(self._font_size)

        # Set font for all script types
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), self._font_name)
        rFonts.set(qn("w:hAnsi"), self._font_name)
        rFonts.set(qn("w:cs"), self._font_name)

        # Set default paragraph spacing
        pPr = style.element.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:line"), "276")  # 1.15 line spacing

    # ═══════════════════════════════════════════════════════════════════════════
    # Page Budget System
    # ═══════════════════════════════════════════════════════════════════════════

    def _calculate_spacing_scale(self) -> None:
        """Calculate spacing scale factor to fit exercises within page budget.

        Accounts for lesson illustrations and solved examples in study book mode.
        Each lesson illustration takes ~1.5 pages, each solved example takes ~1 page.
        """
        if not self.exercises and not self.solved_examples and not self.lesson_illustrations:
            self._spacing_scale = 1.0
            return

        total_height_cm = 0.0
        for ex in self.exercises:
            ex_type = ex.get("type", "show_work")
            total_height_cm += self.EXERCISE_HEIGHT_CM.get(ex_type, 5.0)

        # Section header overhead
        topics = set(ex.get("topic", "General") for ex in self.exercises) if self.exercises else set()
        section_overhead_cm = len(topics) * 2.5
        total_height_cm += section_overhead_cm

        # Account for lesson illustrations (~1.5 pages each)
        illustration_height_cm = len(self.lesson_illustrations) * (self._USABLE_HEIGHT_CM * 1.5)
        total_height_cm += illustration_height_cm

        # Account for solved examples (each takes ~1 page = _USABLE_HEIGHT_CM)
        example_height_cm = len(self.solved_examples) * self._USABLE_HEIGHT_CM
        # Also add section headers for example groups
        example_topics = set(ex.get("topic", "General") for ex in self.solved_examples)
        example_height_cm += len(example_topics) * 2.5

        total_height_cm += example_height_cm

        available_height_cm = self._content_pages * self._USABLE_HEIGHT_CM

        if total_height_cm > 0 and available_height_cm > 0:
            self._spacing_scale = min(1.2, available_height_cm / total_height_cm)
        else:
            self._spacing_scale = 1.0

        # Clamp to reasonable range
        self._spacing_scale = max(0.6, min(1.5, self._spacing_scale))

        logger.info(
            f"[DocxGenerator] Page budget: target={self._target_pages}, "
            f"content_pages={self._content_pages}, exercises={len(self.exercises)}, "
            f"solved_examples={len(self.solved_examples)}, "
            f"lesson_illustrations={len(self.lesson_illustrations)}, "
            f"estimated_height={total_height_cm:.1f}cm, "
            f"available={available_height_cm:.1f}cm, "
            f"spacing_scale={self._spacing_scale:.2f}"
        )

    def _get_answer_lines(self, exercise: dict) -> int:
        """Get number of answer lines for an exercise, respecting page budget."""
        ex_type = exercise.get("type", "show_work")
        difficulty = exercise.get("difficulty", "medium")

        base_lines = {
            "show_work": {"easy": 5, "medium": 7, "hard": 9},
            "long_answer": {"easy": 8, "medium": 10, "hard": 12},
            "word_problem": {"easy": 5, "medium": 7, "hard": 9},
            "word_problems": {"easy": 5, "medium": 7, "hard": 9},
            "fill_in_blank": {"easy": 2, "medium": 2, "hard": 3},
            "fill_blank": {"easy": 2, "medium": 2, "hard": 3},
        }

        type_lines = base_lines.get(ex_type, {"easy": 5, "medium": 7, "hard": 9})
        lines = type_lines.get(difficulty, 7)

        scaled_lines = max(3, int(lines * self._spacing_scale))
        return scaled_lines

    # ═══════════════════════════════════════════════════════════════════════════
    # Cover Page
    # ═══════════════════════════════════════════════════════════════════════════

    def _add_cover_page(self) -> None:
        """Add a clean, minimal cover page with title and student info fields."""
        # Top spacing
        for _ in range(5):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # School/Department name
        if self._school_name:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(self._school_name)
            run.font.name = self._font_name
            run.font.size = Pt(14)
            run.font.color.rgb = Colors.COVER_SUBTITLE
            run.font.bold = True

        # Decorative line
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        run = p.add_run("━" * 40)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x70, 0xAD, 0x47)

        # Main title (large, bold, centered)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(self._title)
        run.font.name = self._font_name
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = Colors.COVER_TITLE

        # Grade and term info
        info_parts = []
        if self._grade:
            info_parts.append(self._grade)
        if self._term:
            info_parts.append(self._term)
        if self._academic_year:
            info_parts.append(self._academic_year)

        if info_parts:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(" — ".join(info_parts))
            run.font.name = self._font_name
            run.font.size = Pt(13)
            run.font.color.rgb = Colors.COVER_SUBTITLE

        # Spacing before student fields
        for _ in range(6):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # Student info fields
        self._add_cover_field("Name")
        self._add_cover_field("Class")
        self._add_cover_field("Date")

        # Page break after cover
        self.doc.add_page_break()

    def _add_cover_field(self, label: str) -> None:
        """Add a student info field with underline space on cover page."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(20)
        p.paragraph_format.left_indent = Cm(3)

        run = p.add_run(f"{label}: ")
        run.font.name = self._font_name
        run.font.size = Pt(13)
        run.font.bold = True

        run = p.add_run("_" * 45)
        run.font.name = self._font_name
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ═══════════════════════════════════════════════════════════════════════════
    # Page Headers (3-column layout with green separator)
    # ═══════════════════════════════════════════════════════════════════════════

    def _add_page_headers(self) -> None:
        """Add 3-column header table to every section: Title | School | Year.
        Below it, a thin green horizontal line separator.
        """
        for section in self.doc.sections:
            header = section.header
            header.is_linked_to_previous = False

            # Clear existing header content
            for p in header.paragraphs:
                p.clear()

            # Remove default empty paragraph
            if header.paragraphs:
                header.paragraphs[0]._p.getparent().remove(header.paragraphs[0]._p)

            # Create a 3-column table for header layout (no visible borders)
            header_table = OxmlElement("w:tbl")

            # Table properties — full width, no borders
            tblPr = OxmlElement("w:tblPr")
            tblW = OxmlElement("w:tblW")
            tblW.set(qn("w:w"), "5000")
            tblW.set(qn("w:type"), "pct")
            tblPr.append(tblW)

            # No borders on header table
            tblBorders = OxmlElement("w:tblBorders")
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border_el = OxmlElement(f"w:{border_name}")
                border_el.set(qn("w:val"), "nil")
                tblBorders.append(border_el)
            tblPr.append(tblBorders)
            header_table.append(tblPr)

            # Table grid — 3 columns
            tblGrid = OxmlElement("w:tblGrid")
            for width in [3000, 4000, 3000]:
                gridCol = OxmlElement("w:gridCol")
                gridCol.set(qn("w:w"), str(width))
                tblGrid.append(gridCol)
            header_table.append(tblGrid)

            # Single row with 3 cells
            tr = OxmlElement("w:tr")

            # Cell 1: Title (left-aligned)
            tc1 = self._create_header_cell(self._title, "left", bold=True)
            tr.append(tc1)

            # Cell 2: School name (center-aligned)
            tc2 = self._create_header_cell(self._school_name, "center", bold=False)
            tr.append(tc2)

            # Cell 3: Year (right-aligned)
            year_text = self._academic_year if self._academic_year else ""
            tc3 = self._create_header_cell(year_text, "right", bold=False)
            tr.append(tc3)

            header_table.append(tr)

            # Add table to header
            header._element.append(header_table)

            # Add green separator line below header table
            sep_p = OxmlElement("w:p")
            sep_pPr = OxmlElement("w:pPr")
            # Bottom border as separator
            pBdr = OxmlElement("w:pBdr")
            bottom_border = OxmlElement("w:bottom")
            bottom_border.set(qn("w:val"), "single")
            bottom_border.set(qn("w:sz"), "12")
            bottom_border.set(qn("w:space"), "1")
            bottom_border.set(qn("w:color"), Colors.SEPARATOR_GREEN)
            pBdr.append(bottom_border)
            sep_pPr.append(pBdr)
            # Small spacing
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:before"), "60")
            spacing.set(qn("w:after"), "60")
            sep_pPr.append(spacing)
            sep_p.append(sep_pPr)
            header._element.append(sep_p)

    def _create_header_cell(self, text: str, alignment: str, bold: bool = False) -> OxmlElement:
        """Create a table cell for the header with specified text and alignment."""
        tc = OxmlElement("w:tc")

        # Cell properties — no borders
        tcPr = OxmlElement("w:tcPr")
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ["top", "left", "bottom", "right"]:
            border_el = OxmlElement(f"w:{border_name}")
            border_el.set(qn("w:val"), "nil")
            tcBorders.append(border_el)
        tcPr.append(tcBorders)
        tc.append(tcPr)

        # Paragraph with text
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")

        # Alignment
        jc = OxmlElement("w:jc")
        align_map = {"left": "left", "center": "center", "right": "right"}
        jc.set(qn("w:val"), align_map.get(alignment, "left"))
        pPr.append(jc)

        # Spacing
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:before"), "0")
        pPr.append(spacing)

        p.append(pPr)

        # Run with text
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), self._font_name)
        rFonts.set(qn("w:hAnsi"), self._font_name)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "18")  # 9pt
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "18")
        rPr.append(szCs)
        if bold:
            b = OxmlElement("w:b")
            rPr.append(b)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "555555")
        rPr.append(color)
        r.append(rPr)

        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p.append(r)

        tc.append(p)
        return tc

    # ═══════════════════════════════════════════════════════════════════════════
    # Page Numbers
    # ═══════════════════════════════════════════════════════════════════════════

    def _add_page_numbers(self) -> None:
        """Add centered page numbers to footer."""
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # "— PAGE —" format
            run = p.add_run("— ")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.name = self._font_name

            # PAGE field
            fldSimple = OxmlElement("w:fldSimple")
            fldSimple.set(qn("w:instr"), " PAGE ")
            fld_run = OxmlElement("w:r")
            fld_rPr = OxmlElement("w:rPr")
            fld_sz = OxmlElement("w:sz")
            fld_sz.set(qn("w:val"), "18")
            fld_rPr.append(fld_sz)
            fld_run.append(fld_rPr)
            fld_text = OxmlElement("w:t")
            fld_text.text = "1"
            fld_run.append(fld_text)
            fldSimple.append(fld_run)
            p._p.append(fldSimple)

            run = p.add_run(" —")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.name = self._font_name

    # ═══════════════════════════════════════════════════════════════════════════
    # Workbook Content — Section Layout
    # ═══════════════════════════════════════════════════════════════════════════

    def _add_workbook_content(self) -> None:
        """Add all exercise sections with page-budget-aware layout.

        In study book mode (when lesson_illustrations/solved_examples are populated), renders:
        - Lesson Illustration FIRST (textbook-style summary)
        - Solved examples BEFORE practice exercises, grouped by topic
        - Structure: [Topic 1 Lesson] → [Topic 1 Examples] → [Topic 1 Exercises] → [Topic 2 Lesson] → ...
        """
        if not self.exercises and not self.solved_examples and not self.lesson_illustrations:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("No exercises generated. Please check your configuration.")
            run.font.size = Pt(12)
            run.font.italic = True
            return

        # Group exercises by topic
        sections = self._group_exercises_by_section()

        # Group solved examples by topic
        examples_by_topic = self._group_solved_examples_by_topic()

        # Group lesson illustrations by topic
        illustrations_by_topic = self._group_lesson_illustrations_by_topic()

        question_num = 1
        example_num = 1
        current_height_cm = 0.0

        # Circled number characters for examples
        circled_numbers = "❶❷❸❹❺❻❼❽❾❿"

        for section_idx, section in enumerate(sections):
            section_title = section.get("title", f"Section {section_idx + 1}")
            num_exercises = len(section.get("exercises", []))

            # ─── Render Lesson Illustration for this topic FIRST ──────────────
            topic_illustration = illustrations_by_topic.get(section_title)
            if topic_illustration:
                # Lesson illustration takes ~1.5 pages
                illustration_height = self._USABLE_HEIGHT_CM * 1.5
                if current_height_cm > 0:
                    # Always start lesson illustration on a new page
                    self.doc.add_page_break()
                    current_height_cm = 0.0

                self._add_lesson_illustration(topic_illustration)
                current_height_cm += illustration_height

                # Page break after lesson illustration
                self.doc.add_page_break()
                current_height_cm = 0.0

            # ─── Render Solved Examples for this topic ────────────────────────
            topic_examples = examples_by_topic.get(section_title, [])
            if topic_examples:
                # Section header for solved examples
                example_header_height = 2.5
                if current_height_cm + example_header_height > self._USABLE_HEIGHT_CM:
                    self.doc.add_page_break()
                    current_height_cm = 0.0

                self._add_examples_section_title(section_title, len(topic_examples))
                current_height_cm += example_header_height

                # Render each solved example
                for ex in topic_examples:
                    # Each solved example takes ~1 page worth of space
                    example_height = 12.0  # Approximate height in cm
                    if current_height_cm + example_height > self._USABLE_HEIGHT_CM:
                        self.doc.add_page_break()
                        current_height_cm = 0.0

                    circled = circled_numbers[example_num - 1] if example_num <= 10 else f"({example_num})"
                    self._add_solved_example_section(ex, circled)
                    example_num += 1
                    current_height_cm += example_height

                # Page break between examples and exercises
                self.doc.add_page_break()
                current_height_cm = 0.0

            # ─── Render Practice Exercises for this topic ────────────────────
            section_header_height = 2.5

            # Check if section header fits
            if current_height_cm + section_header_height > self._USABLE_HEIGHT_CM:
                self.doc.add_page_break()
                current_height_cm = 0.0

            self._add_section_title(section_title, num_exercises)
            current_height_cm += section_header_height

            # Render each exercise in a bordered box
            for ex_idx, exercise in enumerate(section.get("exercises", [])):
                ex_type = exercise.get("type", "show_work")
                ex_height = self.EXERCISE_HEIGHT_CM.get(ex_type, 5.0) * self._spacing_scale

                # Page break if content would overflow
                if current_height_cm + ex_height > self._USABLE_HEIGHT_CM:
                    self.doc.add_page_break()
                    current_height_cm = 0.0

                self._add_exercise_box(exercise, question_num)
                question_num += 1
                current_height_cm += ex_height

            # Between sections: if page is >80% full, break
            if section_idx < len(sections) - 1:
                if current_height_cm > self._USABLE_HEIGHT_CM * 0.80:
                    self.doc.add_page_break()
                    current_height_cm = 0.0

        # Handle any lesson illustrations for topics not in exercises
        remaining_illustration_topics = set(illustrations_by_topic.keys()) - set(s.get("title", "") for s in sections)
        for topic in remaining_illustration_topics:
            illustration = illustrations_by_topic[topic]
            if not illustration:
                continue

            self.doc.add_page_break()
            current_height_cm = 0.0
            self._add_lesson_illustration(illustration)
            current_height_cm += self._USABLE_HEIGHT_CM * 1.5

        # Handle any solved examples for topics not in exercises
        remaining_topics = set(examples_by_topic.keys()) - set(s.get("title", "") for s in sections)
        for topic in remaining_topics:
            topic_examples = examples_by_topic[topic]
            if not topic_examples:
                continue

            self.doc.add_page_break()
            current_height_cm = 0.0

            self._add_examples_section_title(topic, len(topic_examples))
            current_height_cm += 2.5

            for ex in topic_examples:
                if current_height_cm + 12.0 > self._USABLE_HEIGHT_CM:
                    self.doc.add_page_break()
                    current_height_cm = 0.0

                circled = circled_numbers[example_num - 1] if example_num <= 10 else f"({example_num})"
                self._add_solved_example_section(ex, circled)
                example_num += 1
                current_height_cm += 12.0

        # Pad remaining pages if under target
        self._pad_to_target_pages()

    def _group_exercises_by_section(self) -> list[dict]:
        """Group exercises into sections by topic."""
        sections_map: dict[str, dict] = {}

        for exercise in self.exercises:
            topic = exercise.get("topic", "General")
            if topic not in sections_map:
                sections_map[topic] = {
                    "title": topic,
                    "exercises": [],
                }
            sections_map[topic]["exercises"].append(exercise)

        return list(sections_map.values())

    def _group_solved_examples_by_topic(self) -> dict[str, list[dict]]:
        """Group solved examples by their topic field.

        Returns:
            Dict mapping topic name to list of solved example dicts.
        """
        grouped: dict[str, list[dict]] = {}
        for example in self.solved_examples:
            topic = example.get("topic", "General")
            if topic not in grouped:
                grouped[topic] = []
            grouped[topic].append(example)
        return grouped

    def _group_lesson_illustrations_by_topic(self) -> dict[str, dict]:
        """Group lesson illustrations by their topic field.

        Returns:
            Dict mapping topic name to the illustration dict for that topic.
        """
        grouped: dict[str, dict] = {}
        for illustration in self.lesson_illustrations:
            topic = illustration.get("topic", "General")
            grouped[topic] = illustration
        return grouped

    # ═══════════════════════════════════════════════════════════════════════════
    # Lesson Illustration Section (Study Book Mode)
    # ═══════════════════════════════════════════════════════════════════════════

    def _add_lesson_illustration(self, illustration: dict) -> None:
        """Render a full lesson illustration section for a topic.

        Layout:
        - Section Title bar: '📖 Lesson: [Topic Name]' with green background
        - Introduction paragraph
        - Key Concepts box (green border, bullet points)
        - Theorems section (each with name, statement, notation)
        - Key Formulas box (blue border, OMML math rendering)
        - Important Notes box (orange border, yellow background)
        """
        topic = illustration.get("topic", "Lesson")
        introduction = illustration.get("introduction", "")
        key_concepts = illustration.get("key_concepts", [])
        theorems = illustration.get("theorems", [])
        key_formulas = illustration.get("key_formulas", [])
        important_notes = illustration.get("important_notes", [])

        # ── Section Title Bar ──────────────────────────────────────────────
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)

        # Green background shading on the paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), Colors.LESSON_TITLE_BG)
        pPr.append(shd)

        # Title text
        display_title = f"\U0001F4D6 Lesson: {topic}"
        run = p.add_run(display_title)
        run.font.name = self._font_name
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = Colors.LESSON_TITLE_TEXT

        # ── Introduction ───────────────────────────────────────────────────
        if introduction:
            intro_p = self.doc.add_paragraph()
            intro_p.paragraph_format.space_before = Pt(6)
            intro_p.paragraph_format.space_after = Pt(10)
            self._add_text_with_math(intro_p, introduction)

        # ── Key Concepts Box ───────────────────────────────────────────────
        if key_concepts:
            self._add_lesson_concepts_box(key_concepts)

        # ── Theorems Section ───────────────────────────────────────────────
        if theorems:
            self._add_lesson_theorems_section(theorems)

        # ── Key Formulas Box ───────────────────────────────────────────────
        if key_formulas:
            self._add_lesson_formulas_box(key_formulas)

        # ── Important Notes Box ────────────────────────────────────────────
        if important_notes:
            self._add_lesson_notes_box(important_notes)

        # ── Separator after illustration ───────────────────────────────────
        sep_p = self.doc.add_paragraph()
        sep_p.paragraph_format.space_before = Pt(4)
        sep_p.paragraph_format.space_after = Pt(12)
        pPr_sep = sep_p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom_bdr = OxmlElement("w:bottom")
        bottom_bdr.set(qn("w:val"), "single")
        bottom_bdr.set(qn("w:sz"), "6")
        bottom_bdr.set(qn("w:space"), "1")
        bottom_bdr.set(qn("w:color"), Colors.SEPARATOR_GREEN)
        pBdr.append(bottom_bdr)
        pPr_sep.append(pBdr)

    def _add_lesson_concepts_box(self, concepts: list[dict]) -> None:
        """Render the Key Concepts box with green border and bullet points."""
        # Header
        header_p = self.doc.add_paragraph()
        header_p.paragraph_format.space_before = Pt(8)
        header_p.paragraph_format.space_after = Pt(4)
        run = header_p.add_run("📌 Key Concepts")
        run.font.name = self._font_name
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = Colors.LESSON_CONCEPTS_HEADER

        # Concepts in a bordered table (single cell used as a box)
        table = self.doc.add_table(rows=1, cols=1)
        self._set_table_full_width(table)
        cell = table.rows[0].cells[0]

        # Style the cell
        self._set_cell_border(
            cell, top="single", bottom="single", left="single", right="single",
            color=Colors.LESSON_CONCEPTS_BORDER, size="6"
        )
        self._set_cell_shading(cell, Colors.LESSON_CONCEPTS_BG)
        self._set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

        # Clear default paragraph
        cell.paragraphs[0].clear()

        for i, concept in enumerate(concepts):
            # LLM occasionally returns plain strings instead of {name, definition} dicts.
            if isinstance(concept, str):
                concept = {"name": "", "definition": concept}
            elif not isinstance(concept, dict):
                continue
            name = concept.get("name", concept.get("concept", ""))
            definition = concept.get("definition", concept.get("description", ""))

            if i > 0:
                # Add spacing between concepts
                spacer = cell.add_paragraph()
                spacer.paragraph_format.space_before = Pt(2)
                spacer.paragraph_format.space_after = Pt(2)

            # Concept name (bold bullet)
            concept_p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            concept_p.paragraph_format.space_before = Pt(3)
            concept_p.paragraph_format.space_after = Pt(2)

            bullet_run = concept_p.add_run("• ")
            bullet_run.font.name = self._font_name
            bullet_run.font.size = Pt(self._font_size)

            if name:
                name_run = concept_p.add_run(f"{name}: ")
                name_run.font.name = self._font_name
                name_run.font.size = Pt(self._font_size)
                name_run.font.bold = True

            # Definition with math support
            if definition:
                self._add_text_with_math(concept_p, definition)

    def _add_lesson_theorems_section(self, theorems: list[dict]) -> None:
        """Render theorems with name, statement, and optional notation."""
        # Header
        header_p = self.doc.add_paragraph()
        header_p.paragraph_format.space_before = Pt(10)
        header_p.paragraph_format.space_after = Pt(4)
        run = header_p.add_run("📐 Theorems & Properties")
        run.font.name = self._font_name
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = Colors.LESSON_CONCEPTS_HEADER

        for i, theorem in enumerate(theorems):
            if isinstance(theorem, str):
                theorem = {"name": f"Theorem {i + 1}", "statement": theorem}
            elif not isinstance(theorem, dict):
                continue
            name = theorem.get("name", f"Theorem {i + 1}")
            statement = theorem.get("statement", "")
            notation = theorem.get("notation", "")

            # Theorem name
            name_p = self.doc.add_paragraph()
            name_p.paragraph_format.space_before = Pt(6)
            name_p.paragraph_format.space_after = Pt(2)
            name_p.paragraph_format.left_indent = Pt(12)

            name_run = name_p.add_run(f"▸ {name}")
            name_run.font.name = self._font_name
            name_run.font.size = Pt(self._font_size)
            name_run.font.bold = True

            # Statement
            if statement:
                stmt_p = self.doc.add_paragraph()
                stmt_p.paragraph_format.space_before = Pt(2)
                stmt_p.paragraph_format.space_after = Pt(2)
                stmt_p.paragraph_format.left_indent = Pt(24)
                self._add_text_with_math(stmt_p, statement)

            # Notation (math formula)
            if notation:
                notation_p = self.doc.add_paragraph()
                notation_p.paragraph_format.space_before = Pt(2)
                notation_p.paragraph_format.space_after = Pt(4)
                notation_p.paragraph_format.left_indent = Pt(24)

                label_run = notation_p.add_run("Notation: ")
                label_run.font.name = self._font_name
                label_run.font.size = Pt(self._font_size)
                label_run.font.italic = True
                self._add_text_with_math(notation_p, notation)

            # Separator between theorems (except last)
            if i < len(theorems) - 1:
                sep_p = self.doc.add_paragraph()
                sep_p.paragraph_format.space_before = Pt(2)
                sep_p.paragraph_format.space_after = Pt(2)
                pPr = sep_p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom_bdr = OxmlElement("w:bottom")
                bottom_bdr.set(qn("w:val"), "dotted")
                bottom_bdr.set(qn("w:sz"), "4")
                bottom_bdr.set(qn("w:space"), "1")
                bottom_bdr.set(qn("w:color"), Colors.LESSON_THEOREMS_BORDER)
                pBdr.append(bottom_bdr)
                pPr.append(pBdr)

    def _add_lesson_formulas_box(self, formulas: list[dict]) -> None:
        """Render the Key Formulas box with blue border and OMML math."""
        # Header
        header_p = self.doc.add_paragraph()
        header_p.paragraph_format.space_before = Pt(10)
        header_p.paragraph_format.space_after = Pt(4)
        run = header_p.add_run("📝 Key Formulas")
        run.font.name = self._font_name
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = Colors.LESSON_FORMULAS_HEADER

        # Formulas in a bordered table (single cell box)
        table = self.doc.add_table(rows=1, cols=1)
        self._set_table_full_width(table)
        cell = table.rows[0].cells[0]

        # Style the cell
        self._set_cell_border(
            cell, top="single", bottom="single", left="single", right="single",
            color=Colors.LESSON_FORMULAS_BORDER, size="6"
        )
        self._set_cell_shading(cell, Colors.LESSON_FORMULAS_BG)
        self._set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

        # Clear default paragraph
        cell.paragraphs[0].clear()

        for i, formula in enumerate(formulas):
            if isinstance(formula, str):
                formula = {"name": "", "formula": formula}
            elif not isinstance(formula, dict):
                continue
            name = formula.get("name", formula.get("label", ""))
            expression = formula.get("formula", formula.get("expression", ""))
            description = formula.get("description", "")

            formula_p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            formula_p.paragraph_format.space_before = Pt(4)
            formula_p.paragraph_format.space_after = Pt(4)

            # Formula name/label
            if name:
                label_run = formula_p.add_run(f"{name}: ")
                label_run.font.name = self._font_name
                label_run.font.size = Pt(self._font_size)
                label_run.font.bold = True

            # Formula expression (with math rendering)
            if expression:
                self._add_text_with_math(formula_p, expression)

            # Description on next line if present
            if description:
                desc_p = cell.add_paragraph()
                desc_p.paragraph_format.space_before = Pt(1)
                desc_p.paragraph_format.space_after = Pt(3)
                desc_p.paragraph_format.left_indent = Pt(12)
                desc_run = desc_p.add_run(f"  → {description}")
                desc_run.font.name = self._font_name
                desc_run.font.size = Pt(self._font_size - 1)
                desc_run.font.italic = True

    def _add_lesson_notes_box(self, notes: list[str]) -> None:
        """Render the Important Notes box with orange border and yellow background."""
        # Header
        header_p = self.doc.add_paragraph()
        header_p.paragraph_format.space_before = Pt(10)
        header_p.paragraph_format.space_after = Pt(4)
        run = header_p.add_run("⚠️ Important Notes")
        run.font.name = self._font_name
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = Colors.LESSON_NOTES_HEADER

        # Notes in a bordered table (single cell box)
        table = self.doc.add_table(rows=1, cols=1)
        self._set_table_full_width(table)
        cell = table.rows[0].cells[0]

        # Style the cell
        self._set_cell_border(
            cell, top="single", bottom="single", left="single", right="single",
            color=Colors.LESSON_NOTES_BORDER, size="6"
        )
        self._set_cell_shading(cell, Colors.LESSON_NOTES_BG)
        self._set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

        # Clear default paragraph
        cell.paragraphs[0].clear()

        for i, note in enumerate(notes):
            note_text = note if isinstance(note, str) else note.get("text", str(note))

            note_p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            note_p.paragraph_format.space_before = Pt(3)
            note_p.paragraph_format.space_after = Pt(3)

            bullet_run = note_p.add_run("⚡ ")
            bullet_run.font.name = self._font_name
            bullet_run.font.size = Pt(self._font_size)

            self._add_text_with_math(note_p, note_text)

    # ═══════════════════════════════════════════════════════════════════════════
    # Solved Examples Section (Study Book Mode)
    # ═══════════════════════════════════════════════════════════════════════════

    def _add_examples_section_title(self, title: str, example_count: int) -> None:
        """Add section title for solved examples: 'SOLVED EXAMPLES: Topic (N)' in green."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(10)

        # Add light green background shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), Colors.EXAMPLE_SECTION_BG)
        pPr.append(shd)

        # Section title text
        display_title = f"SOLVED EXAMPLES: {title} ({example_count})"
        run = p.add_run(display_title)
        run.font.name = self._font_name
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = Colors.EXAMPLE_TITLE_TEXT
        run.font.underline = True

    def _add_solved_example_section(self, example: dict, circled_number: str) -> None:
        """Render a single solved example with solution box, formula box, and coefficients.

        Layout:
        - Example number with circled number in a small bordered box
        - Problem statement in bold, dark green, underlined
        - Solution box: bordered rectangle with step-by-step solution
        - Key Formula box (optional): pink/salmon background
        - Coefficients box (optional): light blue background

        Args:
            example: Solved example dict with title, solution_steps, key_formula, coefficients.
            circled_number: The circled number character (❶❷❸ etc.)
        """
        title = example.get("title", "")
        steps = example.get("solution_steps", [])
        key_formula = example.get("key_formula")
        coefficients = example.get("coefficients")

        # ─── Example Number + Title ─────────────────────────────────────────
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)

        # Circled number in a bordered box with dark blue background
        run = p.add_run(f" {circled_number} ")
        run.font.name = self._font_name
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Add dark blue background to the number
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), Colors.EXAMPLE_NUMBER_BG)
        rPr.append(shd)
        # Add border around the number
        bdr = OxmlElement("w:bdr")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "8")
        bdr.set(qn("w:space"), "2")
        bdr.set(qn("w:color"), Colors.EXAMPLE_NUMBER_BG)
        rPr.append(bdr)

        # Space between number and title
        run = p.add_run("  ")
        run.font.size = Pt(12)

        # Problem statement (bold, dark green)
        if MathParser.has_math(title):
            self._add_text_with_math(p, title, bold=True)
        else:
            run = p.add_run(title)
            run.font.name = self._font_name
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = Colors.EXAMPLE_TITLE_TEXT
            run.font.underline = True

        # ─── Solution Box (bordered rectangle) ──────────────────────────────
        if steps:
            table = self.doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._set_table_full_width(table)

            cell = table.cell(0, 0)
            self._set_cell_border(
                cell,
                top="single", bottom="single", left="single", right="single",
                color=Colors.EXAMPLE_BOX_BORDER, size="12"
            )
            self._set_cell_shading(cell, Colors.EXAMPLE_BOX_BG)
            self._set_cell_margins(cell, top=150, bottom=150, left=200, right=200)

            # Solution header
            header_p = cell.paragraphs[0]
            header_p.paragraph_format.space_after = Pt(8)
            run = header_p.add_run("Solution")
            run.font.name = self._font_name
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.underline = True
            run.font.color.rgb = RGBColor(0x2F, 0x5F, 0x8F)

            # Each solution step
            for step in steps:
                step_p = cell.add_paragraph()
                step_p.paragraph_format.space_before = Pt(3)
                step_p.paragraph_format.space_after = Pt(3)
                step_p.paragraph_format.left_indent = Cm(0.5)

                # Determine prefix styling
                step_text = str(step).strip()
                is_conclusion = step_text.startswith("\u2234")  # ∴
                is_reasoning = step_text.startswith("\u2235")  # ∵
                is_alternative = step_text.lower().startswith("or ") or step_text.startswith("\u0623\u0648 ")

                if is_conclusion:
                    # Conclusion steps: bold, slightly more spacing
                    step_p.paragraph_format.space_before = Pt(6)
                    self._add_text_with_math(step_p, step_text, bold=True)
                elif is_alternative:
                    # Alternative solutions: indented
                    step_p.paragraph_format.left_indent = Cm(1.0)
                    self._add_text_with_math(step_p, step_text, bold=False)
                elif is_reasoning:
                    # Reasoning steps: normal
                    self._add_text_with_math(step_p, step_text, bold=False)
                else:
                    # Regular step
                    self._add_text_with_math(step_p, step_text, bold=False)

            # Spacer after solution box
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            spacer.paragraph_format.space_before = Pt(0)

        # ─── Key Formula Box (optional, pink/salmon background) ──────────────
        if key_formula:
            table = self.doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._set_table_full_width(table)

            cell = table.cell(0, 0)
            self._set_cell_border(
                cell,
                top="single", bottom="single", left="single", right="single",
                color=Colors.FORMULA_BOX_BORDER, size="8"
            )
            self._set_cell_shading(cell, Colors.FORMULA_BOX_BG)
            self._set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

            # Formula header
            formula_p = cell.paragraphs[0]
            formula_p.paragraph_format.space_after = Pt(4)
            formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = formula_p.add_run("Key Formula: ")
            run.font.name = self._font_name
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x44, 0x44)

            # Formula content with OMML math
            formula_content_p = cell.add_paragraph()
            formula_content_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            formula_content_p.paragraph_format.space_before = Pt(4)
            formula_content_p.paragraph_format.space_after = Pt(4)
            self._add_text_with_math(formula_content_p, str(key_formula), bold=True)

            # Spacer
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)
            spacer.paragraph_format.space_before = Pt(0)

        # ─── Coefficients Box (optional, light blue background) ──────────────
        if coefficients and isinstance(coefficients, dict):
            table = self.doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._set_table_full_width(table)

            cell = table.cell(0, 0)
            self._set_cell_border(
                cell,
                top="single", bottom="single", left="single", right="single",
                color=Colors.COEFFICIENTS_BOX_BORDER, size="6"
            )
            self._set_cell_shading(cell, Colors.COEFFICIENTS_BOX_BG)
            self._set_cell_margins(cell, top=80, bottom=80, left=150, right=150)

            # Coefficients content
            coeff_p = cell.paragraphs[0]
            coeff_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            coeff_p.paragraph_format.space_after = Pt(2)

            run = coeff_p.add_run("Where: ")
            run.font.name = self._font_name
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

            # List coefficients inline
            coeff_parts = []
            for key, value in coefficients.items():
                coeff_parts.append(f"{key} = {value}")
            coeff_text = ",  ".join(coeff_parts)

            run = coeff_p.add_run(coeff_text)
            run.font.name = self._font_name
            run.font.size = Pt(11)
            run.font.bold = True

            # Spacer
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(8)
            spacer.paragraph_format.space_before = Pt(0)

        # Final spacer between examples
        final_spacer = self.doc.add_paragraph()
        final_spacer.paragraph_format.space_after = Pt(12)
        final_spacer.paragraph_format.space_before = Pt(0)

    def _pad_to_target_pages(self) -> None:
        """Add blank pages if document is under target page count.
        Uses empty paragraphs with page breaks to reach target.
        """
        # Estimate current pages based on content
        # This is approximate — Word will finalize pagination
        estimated_pages = 1 if self._include_cover else 0
        total_height = 0.0
        for ex in self.exercises:
            ex_type = ex.get("type", "show_work")
            total_height += self.EXERCISE_HEIGHT_CM.get(ex_type, 5.0) * self._spacing_scale

        topics = set(ex.get("topic", "General") for ex in self.exercises)
        total_height += len(topics) * 2.5

        estimated_pages += max(1, int(total_height / self._USABLE_HEIGHT_CM) + 1)

        pages_to_add = self._target_pages - estimated_pages
        if pages_to_add > 0:
            for _ in range(pages_to_add):
                self.doc.add_page_break()
                p = self.doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)

    # ═══════════════════════════════════════════════════════════════════════════
    # Section Title (highlighted, purple text)
    # ═══════════════════════════════════════════════════════════════════════════

    def _add_section_title(self, title: str, exercise_count: int) -> None:
        """Add section title: 'PRACTICE (N)' in bold purple with green highlight."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(10)

        # Add highlight/shading to the paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), Colors.SECTION_TITLE_BG)
        pPr.append(shd)

        # Section title text
        display_title = f"PRACTICE: {title} ({exercise_count})"
        run = p.add_run(display_title)
        run.font.name = self._font_name
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = Colors.SECTION_TITLE_TEXT
        run.font.underline = True

    # ═══════════════════════════════════════════════════════════════════════════
    # Exercise Box Rendering (thick black border, beige background)
    # ═══════════════════════════════════════════════════════════════════════════

    def _add_exercise_box(self, exercise: dict, question_num: int) -> None:
        """Add a single exercise in a thick-bordered box with beige background.
        Dispatches to the appropriate renderer based on exercise type.
        """
        exercise_type = exercise.get("type", "show_work")

        if exercise_type == "multiple_choice":
            self._add_mcq_box(exercise, question_num)
        elif exercise_type in ("fill_in_blank", "fill_blank"):
            self._add_fill_blank_box(exercise, question_num)
        elif exercise_type == "long_answer":
            self._add_long_answer_box(exercise, question_num)
        elif exercise_type == "true_false":
            self._add_true_false_box(exercise, question_num)
        elif exercise_type == "matching":
            self._add_matching_box(exercise, question_num)
        else:
            # show_work, word_problem, etc.
            self._add_long_answer_box(exercise, question_num)

    # ─── MCQ Box ──────────────────────────────────────────────────────────────

    def _add_mcq_box(self, exercise: dict, question_num: int) -> None:
        """Add a multiple-choice question in a thick-bordered box.

        Layout:
        - Question text at top (bold)
        - Each option with a bordered letter box + option text
        - Beige background
        """
        # Create a 1-cell table for the question box
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set table width to full page
        self._set_table_full_width(table)

        cell = table.cell(0, 0)

        # Set thick black border on the cell
        self._set_cell_border(
            cell,
            top="single", bottom="single", left="single", right="single",
            color=Colors.QUESTION_BOX_BORDER, size="18"
        )

        # Set beige background
        self._set_cell_shading(cell, Colors.QUESTION_BOX_BG)

        # Cell margins for padding
        self._set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

        # Question text paragraph
        question_text = exercise.get("question", exercise.get("statement", exercise.get("text", "")))
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"Q{question_num}: ")
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size)
        run.font.bold = True

        # Add question text with math support
        self._add_text_with_math(p, question_text, bold=True)

        # Options
        options = exercise.get("options", exercise.get("choices", []))
        option_labels = self._ENGLISH_OPTIONS

        for i, option in enumerate(options):
            if i >= len(option_labels):
                break

            # Clean option text
            option_text = self._clean_option_text(str(option))

            # Create option paragraph
            opt_p = cell.add_paragraph()
            opt_p.paragraph_format.space_before = Pt(4)
            opt_p.paragraph_format.space_after = Pt(4)
            opt_p.paragraph_format.left_indent = Cm(0.5)

            # Option letter in a bordered box style: [A]
            run = opt_p.add_run(f"  {option_labels[i]}  ")
            run.font.name = self._font_name
            run.font.size = Pt(self._font_size)
            run.font.bold = True
            # Add border effect using a box character approach
            rPr = run._r.get_or_add_rPr()
            bdr = OxmlElement("w:bdr")
            bdr.set(qn("w:val"), "single")
            bdr.set(qn("w:sz"), "6")
            bdr.set(qn("w:space"), "1")
            bdr.set(qn("w:color"), Colors.OPTION_LETTER_BORDER)
            rPr.append(bdr)

            # Space between letter and option text
            run = opt_p.add_run("   ")
            run.font.name = self._font_name
            run.font.size = Pt(self._font_size)

            # Option text with math support
            self._add_text_with_math(opt_p, option_text, bold=False)

        # Add spacing paragraph after the table
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        spacer.paragraph_format.space_before = Pt(0)

    # ─── Fill-in-the-Blank Box ────────────────────────────────────────────────

    def _add_fill_blank_box(self, exercise: dict, question_num: int) -> None:
        """Add a fill-in-the-blank question in a thick-bordered box.

        Layout:
        - Question text with visible blank line
        - Answer space below
        """
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_full_width(table)

        cell = table.cell(0, 0)
        self._set_cell_border(
            cell,
            top="single", bottom="single", left="single", right="single",
            color=Colors.QUESTION_BOX_BORDER, size="18"
        )
        self._set_cell_shading(cell, Colors.QUESTION_BOX_BG)
        self._set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

        # Question text
        question_text = exercise.get("question", exercise.get("statement", exercise.get("text", "")))
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(f"Q{question_num}: ")
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size)
        run.font.bold = True

        self._add_text_with_math(p, question_text, bold=True)

        # Blank line for answer
        ans_p = cell.add_paragraph()
        ans_p.paragraph_format.space_before = Pt(8)
        ans_p.paragraph_format.space_after = Pt(4)
        ans_p.paragraph_format.left_indent = Cm(0.5)
        run = ans_p.add_run("Answer: ")
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size)
        run.font.bold = True

        run = ans_p.add_run("_" * 40)
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Additional answer lines if needed
        num_lines = self._get_answer_lines(exercise)
        for _ in range(min(num_lines, 3)):
            line_p = cell.add_paragraph()
            line_p.paragraph_format.space_before = Pt(6)
            line_p.paragraph_format.space_after = Pt(2)
            run = line_p.add_run("_" * 60)
            run.font.name = self._font_name
            run.font.size = Pt(self._font_size)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # Spacer after box
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        spacer.paragraph_format.space_before = Pt(0)

    # ─── Long Answer / Show Work Box ─────────────────────────────────────────

    def _add_long_answer_box(self, exercise: dict, question_num: int) -> None:
        """Add a long-answer/proof question in a thick-bordered box.

        Layout:
        - Question text
        - Tip section with ※ icon in light blue sub-box (if hint provided)
        - Large blank space for student work
        """
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_full_width(table)

        cell = table.cell(0, 0)
        self._set_cell_border(
            cell,
            top="single", bottom="single", left="single", right="single",
            color=Colors.QUESTION_BOX_BORDER, size="18"
        )
        self._set_cell_shading(cell, Colors.QUESTION_BOX_BG)
        self._set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

        # Question text
        question_text = exercise.get("question", exercise.get("statement", exercise.get("text", "")))
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"Q{question_num}: ")
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size)
        run.font.bold = True

        self._add_text_with_math(p, question_text, bold=True)

        # Tip/Hint section (light blue sub-box)
        hint = exercise.get("hint", "")
        if hint:
            self._add_tip_section(cell, hint)

        # Blank space for student work (ruled lines)
        num_lines = self._get_answer_lines(exercise)
        for i in range(num_lines):
            line_p = cell.add_paragraph()
            line_p.paragraph_format.space_before = Pt(0)
            line_p.paragraph_format.space_after = Pt(0)

            # Add bottom border to simulate ruled line
            pPr = line_p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)

            # Set line height for writing space
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            spacing.set(qn("w:line"), "400")  # ~20pt line height
            spacing.set(qn("w:lineRule"), "exact")
            pPr.append(spacing)

            # Empty run to maintain line
            run = line_p.add_run(" ")
            run.font.size = Pt(self._font_size)

        # Spacer after box
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        spacer.paragraph_format.space_before = Pt(0)

    def _add_tip_section(self, cell, hint_text: str) -> None:
        """Add a tip section with ※ icon in a light blue background within the cell."""
        # Tip paragraph with shading
        tip_p = cell.add_paragraph()
        tip_p.paragraph_format.space_before = Pt(6)
        tip_p.paragraph_format.space_after = Pt(8)
        tip_p.paragraph_format.left_indent = Cm(0.3)
        tip_p.paragraph_format.right_indent = Cm(0.3)

        # Add light blue background shading to the paragraph
        pPr = tip_p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), Colors.TIP_BOX_BG)
        pPr.append(shd)

        # Add left border for visual distinction
        pBdr = OxmlElement("w:pBdr")
        left_border = OxmlElement("w:left")
        left_border.set(qn("w:val"), "single")
        left_border.set(qn("w:sz"), "12")
        left_border.set(qn("w:space"), "4")
        left_border.set(qn("w:color"), "4472C4")
        pBdr.append(left_border)
        pPr.append(pBdr)

        # ※ Tip icon and text
        run = tip_p.add_run("※ Tip: ")
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size - 1)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2F, 0x5F, 0x8F)

        run = tip_p.add_run(hint_text)
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size - 1)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ─── True/False Box ───────────────────────────────────────────────────────

    def _add_true_false_box(self, exercise: dict, question_num: int) -> None:
        """Add a true/false question in a bordered box."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_full_width(table)

        cell = table.cell(0, 0)
        self._set_cell_border(
            cell,
            top="single", bottom="single", left="single", right="single",
            color=Colors.QUESTION_BOX_BORDER, size="18"
        )
        self._set_cell_shading(cell, Colors.QUESTION_BOX_BG)
        self._set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

        # Question text
        question_text = exercise.get("question", exercise.get("statement", exercise.get("text", "")))
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(f"Q{question_num}: ")
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size)
        run.font.bold = True

        self._add_text_with_math(p, question_text, bold=True)

        # True / False options with bordered boxes
        opt_p = cell.add_paragraph()
        opt_p.paragraph_format.space_before = Pt(6)
        opt_p.paragraph_format.left_indent = Cm(1.0)

        # True option
        run = opt_p.add_run("  True  ")
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size)
        run.font.bold = True
        rPr = run._r.get_or_add_rPr()
        bdr = OxmlElement("w:bdr")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "6")
        bdr.set(qn("w:space"), "1")
        bdr.set(qn("w:color"), "000000")
        rPr.append(bdr)

        run = opt_p.add_run("          ")
        run.font.size = Pt(self._font_size)

        # False option
        run = opt_p.add_run("  False  ")
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size)
        run.font.bold = True
        rPr = run._r.get_or_add_rPr()
        bdr = OxmlElement("w:bdr")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "6")
        bdr.set(qn("w:space"), "1")
        bdr.set(qn("w:color"), "000000")
        rPr.append(bdr)

        # Spacer
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        spacer.paragraph_format.space_before = Pt(0)

    # ─── Matching Box ─────────────────────────────────────────────────────────

    def _add_matching_box(self, exercise: dict, question_num: int) -> None:
        """Add a matching exercise in a bordered box with two columns."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_full_width(table)

        cell = table.cell(0, 0)
        self._set_cell_border(
            cell,
            top="single", bottom="single", left="single", right="single",
            color=Colors.QUESTION_BOX_BORDER, size="18"
        )
        self._set_cell_shading(cell, Colors.QUESTION_BOX_BG)
        self._set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

        # Question text
        question_text = exercise.get("question", exercise.get("statement", "Match the following:"))
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(f"Q{question_num}: ")
        run.font.name = self._font_name
        run.font.size = Pt(self._font_size)
        run.font.bold = True

        self._add_text_with_math(p, question_text, bold=True)

        # Get matching items
        left_items = exercise.get("left_items", exercise.get("column_a", []))
        right_items = exercise.get("right_items", exercise.get("column_b", []))

        if left_items and right_items:
            num_rows = max(len(left_items), len(right_items))

            # Create matching table inside the cell using paragraphs
            for i in range(num_rows):
                match_p = cell.add_paragraph()
                match_p.paragraph_format.space_before = Pt(3)
                match_p.paragraph_format.space_after = Pt(3)
                match_p.paragraph_format.left_indent = Cm(0.5)

                left_text = left_items[i] if i < len(left_items) else ""
                right_text = right_items[i] if i < len(right_items) else ""
                right_label = chr(65 + i) if i < 26 else str(i + 1)

                run = match_p.add_run(f"{i + 1}. {left_text}")
                run.font.name = self._font_name
                run.font.size = Pt(self._font_size)

                run = match_p.add_run("     ____     ")
                run.font.name = self._font_name
                run.font.size = Pt(self._font_size)

                run = match_p.add_run(f"{right_label}. {right_text}")
                run.font.name = self._font_name
                run.font.size = Pt(self._font_size)

        # Spacer
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        spacer.paragraph_format.space_before = Pt(0)

    # ═══════════════════════════════════════════════════════════════════════════
    # Math Text Rendering (inline OMML or plain text with Unicode)
    # ═══════════════════════════════════════════════════════════════════════════

    def _add_text_with_math(self, paragraph, text: str, bold: bool = False) -> None:
        """Add text to a paragraph, rendering math expressions as OMML when detected.

        For simple text or Unicode math symbols, adds as regular runs.
        For complex math (superscripts, fractions, radicals), inserts OMML elements.
        """
        if not text:
            return

        if MathParser.has_math(text):
            # Split text into math and non-math segments
            self._render_mixed_math_text(paragraph, text, bold)
        else:
            # Plain text — just add as a run
            run = paragraph.add_run(text)
            run.font.name = self._font_name
            run.font.size = Pt(self._font_size)
            if bold:
                run.font.bold = True

    def _render_mixed_math_text(self, paragraph, text: str, bold: bool = False) -> None:
        """Render text that contains a mix of plain text and math expressions.

        Identifies math segments and renders them as OMML, with plain text as regular runs.
        """
        # Find all math patterns in the text
        patterns = [
            (MathParser.FRACTION_RE, "frac"),
            (MathParser.SQRT_RE, "sqrt"),
            (MathParser.SUPERSCRIPT_RE, "sup"),
            (MathParser.SUBSCRIPT_RE, "sub"),
        ]

        # Collect all matches with their positions
        all_matches = []
        for pattern, ptype in patterns:
            for match in pattern.finditer(text):
                all_matches.append((match.start(), match.end(), match.group(0), ptype))

        if not all_matches:
            # No matches found — add as plain text
            run = paragraph.add_run(text)
            run.font.name = self._font_name
            run.font.size = Pt(self._font_size)
            if bold:
                run.font.bold = True
            return

        # Sort by position
        all_matches.sort(key=lambda x: x[0])

        # Remove overlapping matches (keep earliest)
        filtered_matches = []
        last_end = 0
        for start, end, match_text, ptype in all_matches:
            if start >= last_end:
                filtered_matches.append((start, end, match_text, ptype))
                last_end = end

        # Render segments
        pos = 0
        for start, end, match_text, ptype in filtered_matches:
            # Plain text before this match
            if start > pos:
                plain = text[pos:start]
                run = paragraph.add_run(plain)
                run.font.name = self._font_name
                run.font.size = Pt(self._font_size)
                if bold:
                    run.font.bold = True

            # Math segment — insert OMML
            oMath = MathParser.parse_to_omml(match_text)
            paragraph._p.append(oMath)

            pos = end

        # Remaining plain text after last match
        if pos < len(text):
            remaining = text[pos:]
            run = paragraph.add_run(remaining)
            run.font.name = self._font_name
            run.font.size = Pt(self._font_size)
            if bold:
                run.font.bold = True

    # ═══════════════════════════════════════════════════════════════════════════
    # Table & Cell Utility Methods
    # ═══════════════════════════════════════════════════════════════════════════

    def _set_table_full_width(self, table) -> None:
        """Set a table to occupy full page width."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Remove existing width
        existing_w = tblPr.find(qn("w:tblW"))
        if existing_w is not None:
            tblPr.remove(existing_w)

        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")
        tblPr.append(tblW)

        # Set column width
        tblGrid = tbl.find(qn("w:tblGrid"))
        if tblGrid is not None:
            for gridCol in tblGrid.findall(qn("w:gridCol")):
                gridCol.set(qn("w:w"), "9638")  # ~17cm in DXA (full width minus margins)

    def _set_cell_border(
        self,
        cell,
        top: str = "nil",
        bottom: str = "nil",
        left: str = "nil",
        right: str = "nil",
        color: str = "000000",
        size: str = "4",
    ) -> None:
        """Set borders on a table cell.

        Args:
            cell: The table cell to modify.
            top/bottom/left/right: Border style ('nil', 'single', 'double', 'thick').
            color: Hex color string (without #).
            size: Border size in eighths of a point (18 = ~2.25pt thick).
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Remove existing borders
        existing = tcPr.find(qn("w:tcBorders"))
        if existing is not None:
            tcPr.remove(existing)

        tcBorders = OxmlElement("w:tcBorders")

        for edge, style in [
            ("w:top", top),
            ("w:bottom", bottom),
            ("w:left", left),
            ("w:right", right),
        ]:
            element = OxmlElement(edge)
            element.set(qn("w:val"), style)
            if style != "nil":
                element.set(qn("w:sz"), size)
                element.set(qn("w:space"), "0")
                element.set(qn("w:color"), color)
            tcBorders.append(element)

        tcPr.append(tcBorders)

    def _set_cell_shading(self, cell, fill_color: str) -> None:
        """Set background shading on a table cell.

        Args:
            cell: The table cell.
            fill_color: Hex color string (without #).
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Remove existing shading
        existing = tcPr.find(qn("w:shd"))
        if existing is not None:
            tcPr.remove(existing)

        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_color)
        tcPr.append(shd)

    def _set_cell_margins(self, cell, top: int = 0, bottom: int = 0, left: int = 0, right: int = 0) -> None:
        """Set cell margins/padding in DXA units (1440 DXA = 1 inch).

        Args:
            cell: The table cell.
            top/bottom/left/right: Margin in DXA units.
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Remove existing margins
        existing = tcPr.find(qn("w:tcMar"))
        if existing is not None:
            tcPr.remove(existing)

        tcMar = OxmlElement("w:tcMar")

        for edge, value in [("w:top", top), ("w:bottom", bottom), ("w:start", left), ("w:end", right)]:
            margin_el = OxmlElement(edge)
            margin_el.set(qn("w:w"), str(value))
            margin_el.set(qn("w:type"), "dxa")
            tcMar.append(margin_el)

        tcPr.append(tcMar)

    # ═══════════════════════════════════════════════════════════════════════════
    # Text Utility Methods
    # ═══════════════════════════════════════════════════════════════════════════

    def _clean_option_text(self, text: str) -> str:
        """Remove existing option prefixes (A), B., etc.) from option text."""
        prefixes = [
            "A)", "B)", "C)", "D)", "E)", "F)",
            "A.", "B.", "C.", "D.", "E.", "F.",
            "a)", "b)", "c)", "d)", "e)", "f)",
            "a.", "b.", "c.", "d.", "e.", "f.",
        ]
        stripped = text.strip()
        for prefix in prefixes:
            if stripped.startswith(prefix):
                return stripped[len(prefix):].strip()
        return stripped
