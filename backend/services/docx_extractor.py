"""DOCX text extraction service for Word documents."""

import logging
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument

from services.pdf_extractor import PageContent

logger = logging.getLogger(__name__)


class DocxExtractor:
    """Extracts text content from .docx Word documents."""

    # Approximate characters per "page" for pagination simulation
    CHARS_PER_PAGE: int = 3000

    def __init__(self, file_path: str) -> None:
        """Initialize extractor with file path.

        Args:
            file_path: Path to the .docx file.
        """
        self.file_path = Path(file_path)
        self.pages: list[PageContent] = []
        self._total_pages: int = 0

    async def extract(self) -> list[PageContent]:
        """Extract text from a Word document, simulating page boundaries.

        Returns:
            List of PageContent objects (one per simulated page).
        """
        if not self.file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {self.file_path}")

        self.pages = []

        try:
            doc = DocxDocument(str(self.file_path))
        except Exception as e:
            logger.error(f"Failed to open DOCX file: {e}")
            raise RuntimeError(f"Cannot open DOCX file: {self.file_path}") from e

        current_text = ""
        page_num = 1

        # Extract text from paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                current_text += "\n"
                continue

            current_text += para_text + "\n"

            # Check for page break in paragraph runs
            if self._has_page_break(para) or len(current_text) >= self.CHARS_PER_PAGE:
                if current_text.strip():
                    self.pages.append(
                        PageContent(
                            page_num=page_num,
                            text=current_text.strip(),
                            is_ocr=False,
                        )
                    )
                    page_num += 1
                current_text = ""

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    current_text += " | ".join(row_texts) + "\n"

                    if len(current_text) >= self.CHARS_PER_PAGE:
                        if current_text.strip():
                            self.pages.append(
                                PageContent(
                                    page_num=page_num,
                                    text=current_text.strip(),
                                    is_ocr=False,
                                )
                            )
                            page_num += 1
                        current_text = ""

        # Don't forget remaining text
        if current_text.strip():
            self.pages.append(
                PageContent(
                    page_num=page_num,
                    text=current_text.strip(),
                    is_ocr=False,
                )
            )

        self._total_pages = len(self.pages)

        logger.info(
            f"Extracted {self._total_pages} pages from DOCX: {self.file_path.name}"
        )

        return self.pages

    def _has_page_break(self, paragraph) -> bool:
        """Check if a paragraph contains a page break.

        Args:
            paragraph: A python-docx Paragraph object.

        Returns:
            True if the paragraph contains a page break element.
        """
        try:
            from docx.oxml.ns import qn

            for run in paragraph.runs:
                # Check for break elements in the run's XML
                for br in run._r.findall(qn("w:br")):
                    br_type = br.get(qn("w:type"), "")
                    if br_type == "page":
                        return True

            # Check paragraph properties for page break before
            pPr = paragraph._p.find(qn("w:pPr"))
            if pPr is not None:
                # Check for pageBreakBefore
                pb_before = pPr.find(qn("w:pageBreakBefore"))
                if pb_before is not None:
                    return True
        except Exception:
            pass

        return False

    def get_total_pages(self) -> int:
        """Return total page count (simulated from content length).

        Returns:
            Total number of simulated pages.
        """
        if self._total_pages == 0 and self.file_path.exists():
            try:
                doc = DocxDocument(str(self.file_path))
                total_chars = sum(len(p.text) for p in doc.paragraphs)
                self._total_pages = max(1, total_chars // self.CHARS_PER_PAGE + 1)
            except Exception:
                self._total_pages = 1
        return self._total_pages
