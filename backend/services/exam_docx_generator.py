"""Exam DOCX generator — produces professional exam papers matching MOE format.

Generates exam papers with:
  • Official MOE-style header (school name, subject, grade, term, date, time)
  • Organized question sections with mark allocations
  • Optional topic-organized layout (weekly_assessment) + parallel groups
  • RTL Arabic support
  • Answer key with marking rubric
"""

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.section import WD_ORIENT

logger = logging.getLogger(__name__)


class ExamDocxGenerator:
    """Generates exam DOCX files in official MOE format."""

    def __init__(
        self,
        formatting: dict,
        structure: dict,
        letterhead_lines: Optional[list[str]] = None,
        subject_label: Optional[str] = None,
    ) -> None:
        """
        Args:
            formatting: Exam formatting dict (title, school_name, subject, …).
            structure: Exam structure dict (exam_type, total_marks, …).
            letterhead_lines: Optional per-subject ministry letterhead lines.
                              When None, falls back to the math department
                              default for backwards compat with pre-Phase-2
                              callers.
            subject_label: Optional display label for the subject. When set,
                           used as the default `subject` field on the exam
                           header if the formatting dict didn't supply one.
        """
        self.formatting = formatting
        self.structure = structure
        self.language = formatting.get("language", "arabic")
        self.is_rtl = self.language in ("arabic", "bilingual")
        self.exam_type = structure.get("exam_type", "monthly_exam")
        self.is_weekly = self.exam_type == "weekly_assessment"
        self.letterhead_lines = list(letterhead_lines or [])
        self.subject_label = subject_label or ""

    # ─── Public API ──────────────────────────────────────────────────────────

    def generate_exam(
        self,
        exam_content: list[dict],
        output_path: str,
        variant_num: int = 1,
    ) -> str:
        """Generate the exam paper DOCX."""
        doc = Document()
        self._setup_document(doc)
        self._add_exam_header(doc, variant_num)
        self._add_instructions(doc)

        for section in exam_content:
            if not (
                section.get("groups")
                or section.get("questions")
            ):
                continue
            self._add_section(doc, section)

        self._add_exam_footer(doc)
        doc.save(output_path)
        logger.info(f"[ExamDocx] Exam saved: {output_path}")
        return output_path

    def generate_answer_key(
        self, answer_key: list[dict], output_path: str
    ) -> str:
        """Generate the answer key DOCX with marking rubric."""
        doc = Document()
        self._setup_document(doc)
        self._add_answer_key_header(doc)

        for section in answer_key:
            has_answers = (
                section.get("answers") or
                any(g.get("answers") for g in section.get("groups", []))
            )
            if not has_answers:
                continue
            self._add_answer_section(doc, section)

        self._add_marking_rubric(doc, answer_key)
        doc.save(output_path)
        logger.info(f"[ExamDocx] Answer key saved: {output_path}")
        return output_path

    # ─── Document Setup ──────────────────────────────────────────────────────

    def _setup_document(self, doc: Document) -> None:
        """Configure document margins, fonts, and RTL settings."""
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(12)

        if self.is_rtl:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._set_rtl_style(style)

    def _set_rtl_style(self, style) -> None:
        pPr = style.paragraph_format._element
        if pPr is None:
            pPr = OxmlElement("w:pPr")
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)

    def _set_paragraph_rtl(self, paragraph) -> None:
        if self.is_rtl:
            pPr = paragraph._p.get_or_add_pPr()
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            pPr.append(bidi)

    # ─── Exam Header ─────────────────────────────────────────────────────────

    def _add_exam_header(self, doc: Document, variant_num: int = 1) -> None:
        """Add the official MOE-style exam header.

        For weekly_assessment exams, also prepends the multi-line ministry
        letterhead seen on real MOE weekly assessment PDFs.
        """
        if self.is_weekly:
            self._add_ministry_letterhead(doc)

        table = doc.add_table(rows=4, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            row.cells[0].width = Cm(6)
            row.cells[1].width = Cm(5)
            row.cells[2].width = Cm(6)

        school_name = self.formatting.get("school_name", "..................")
        subject = (
            self.formatting.get("subject")
            or self.subject_label
            or ("الرياضيات" if self.is_rtl else "Mathematics")
        )
        grade = self.formatting.get("grade", "")
        term = self.formatting.get("term", "")
        academic_year = self.formatting.get("academic_year", "2025-2026")
        exam_date = self.formatting.get("exam_date", "    /    /      ")
        duration = self.structure.get("duration_minutes", 60)
        total_marks = self.structure.get("total_marks", 40)

        # Title based on exam type
        if self.exam_type == "quiz":
            exam_title = "اختبار قصير" if self.is_rtl else "Quiz"
        elif self.exam_type == "weekly_assessment":
            exam_title = "تقييم أسبوعي" if self.is_rtl else "Weekly Assessment"
        else:
            exam_title = self.formatting.get("title", "امتحان شهري")

        if variant_num > 1:
            exam_title += (
                f" (النموذج {variant_num})" if self.is_rtl
                else f" (Variant {variant_num})"
            )

        if self.is_rtl:
            self._set_cell_text(
                table.cell(0, 2), f"المدرسة: {school_name}", bold=True, size=11
            )
            self._set_cell_text(
                table.cell(0, 1), exam_title, bold=True, size=14, center=True
            )
            self._set_cell_text(
                table.cell(0, 0), "وزارة التربية والتعليم", bold=True, size=11
            )
        else:
            self._set_cell_text(
                table.cell(0, 0), f"School: {school_name}", bold=True, size=11
            )
            self._set_cell_text(
                table.cell(0, 1), exam_title, bold=True, size=14, center=True
            )
            self._set_cell_text(
                table.cell(0, 2), "Ministry of Education", bold=True, size=11
            )

        if self.is_rtl:
            self._set_cell_text(table.cell(1, 2), f"المادة: {subject}", size=11)
            self._set_cell_text(
                table.cell(1, 1), f"الصف: {grade}", size=11, center=True
            )
            self._set_cell_text(
                table.cell(1, 0), f"الفصل الدراسي: {term}", size=11
            )
        else:
            self._set_cell_text(table.cell(1, 0), f"Subject: {subject}", size=11)
            self._set_cell_text(
                table.cell(1, 1), f"Grade: {grade}", size=11, center=True
            )
            self._set_cell_text(table.cell(1, 2), f"Term: {term}", size=11)

        if self.is_rtl:
            self._set_cell_text(
                table.cell(2, 2), f"التاريخ: {exam_date}", size=11
            )
            self._set_cell_text(
                table.cell(2, 1),
                f"العام الدراسي: {academic_year}", size=11, center=True,
            )
            self._set_cell_text(
                table.cell(2, 0), f"الزمن: {duration} دقيقة", size=11
            )
        else:
            self._set_cell_text(table.cell(2, 0), f"Date: {exam_date}", size=11)
            self._set_cell_text(
                table.cell(2, 1),
                f"Academic Year: {academic_year}", size=11, center=True,
            )
            self._set_cell_text(
                table.cell(2, 2), f"Duration: {duration} min", size=11
            )

        if self.is_rtl:
            self._set_cell_text(
                table.cell(3, 2),
                "اسم الطالب: ................................", size=11,
            )
            self._set_cell_text(
                table.cell(3, 1),
                f"الدرجة الكلية: {total_marks}",
                bold=True, size=12, center=True,
            )
            self._set_cell_text(
                table.cell(3, 0), "رقم الجلوس: ............", size=11
            )
        else:
            self._set_cell_text(
                table.cell(3, 0),
                "Student Name: ................................", size=11,
            )
            self._set_cell_text(
                table.cell(3, 1), f"Total Marks: {total_marks}",
                bold=True, size=12, center=True,
            )
            self._set_cell_text(
                table.cell(3, 2), "Seat No: ............", size=11
            )

        self._style_header_table(table)

        # Separator line
        doc.add_paragraph()
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = separator.add_run("═" * 60)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 128)

    def _add_ministry_letterhead(self, doc: Document) -> None:
        """Prepend the multi-line MOE letterhead seen on weekly assessments.

        Lines come from the per-subject SubjectStrategy (passed in via
        constructor). Falls back to the math department default for any
        legacy caller that constructs ExamDocxGenerator without providing
        letterhead_lines.
        """
        if self.letterhead_lines:
            lines = list(self.letterhead_lines)
        else:
            lines = (
                [
                    "وزارة التربية والتعليم و التعليم الفني",
                    "الإدارة المركزية للتعليم العام",
                    "إدارة تنمية مادة الرياضيات",
                    "مكتب مستشار الرياضيات",
                ] if self.is_rtl else [
                    "Ministry of Education & Technical Education",
                    "Central Administration for General Education",
                    "Mathematics Curriculum Development Department",
                    "Office of the Mathematics Consultant",
                ]
            )
        for line in lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_paragraph_rtl(p)
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
        doc.add_paragraph()

    def _add_instructions(self, doc: Document) -> None:
        """Add exam instructions paragraph.

        When parallel groups are used (MOE weekly assessment style), groups
        are alternatives — students answer ONE group only, not all of them.
        The instruction line is adapted accordingly so total marks stay
        consistent with the marks rendered in section headers.
        """
        total_marks = self.structure.get("total_marks", 40)
        duration = self.structure.get("duration_minutes", 60)
        groups = max(1, int(self.structure.get("groups_per_variant", 1) or 1))
        multi_group = groups > 1

        if self.is_rtl:
            if multi_group:
                instructions = (
                    f"أجب عن أسئلة مجموعة واحدة فقط في كل سؤال — "
                    f"الدرجة الكلية ({total_marks}) درجة — "
                    f"الزمن ({duration}) دقيقة"
                )
            else:
                instructions = (
                    f"أجب عن جميع الأسئلة التالية — "
                    f"الدرجة الكلية ({total_marks}) درجة — "
                    f"الزمن ({duration}) دقيقة"
                )
        else:
            if multi_group:
                instructions = (
                    f"Answer ONE group from each section — "
                    f"Total marks: {total_marks} — Time: {duration} minutes"
                )
            else:
                instructions = (
                    f"Answer ALL the following questions — "
                    f"Total marks: {total_marks} — Time: {duration} minutes"
                )

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_rtl(para)
        run = para.add_run(instructions)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(128, 0, 0)
        doc.add_paragraph()

    # ─── Question Sections ───────────────────────────────────────────────────

    def _add_section(self, doc: Document, section: dict) -> None:
        """Add a question section to the exam.

        Renders:
          • Section title with mark allocation (e.g. "First: Algebra [8 marks]")
          • All groups under that section (Group 1, Group 2, …) when more
            than one group is present (MOE weekly-assessment style).
        """
        title = section.get("title_ar" if self.is_rtl else "title_en", "")
        total_marks = section.get("total_marks", 0)

        if self.is_rtl:
            header_text = f"{title}    [{total_marks} درجة]"
        else:
            header_text = f"{title}    [{total_marks} marks]"

        header_para = doc.add_paragraph()
        self._set_paragraph_rtl(header_para)
        header_para.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT if self.is_rtl else WD_ALIGN_PARAGRAPH.LEFT
        )
        run = header_para.add_run(header_text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)

        underline_para = doc.add_paragraph()
        underline_run = underline_para.add_run("─" * 70)
        underline_run.font.size = Pt(8)
        underline_run.font.color.rgb = RGBColor(150, 150, 150)

        question_type = section.get("question_type", "")
        marks_per_q = section.get("marks_per_question", 2)

        groups = section.get("groups")
        if not groups:
            # Backward-compat path: section has a flat questions list
            questions = section.get("questions", [])
            for i, question in enumerate(questions, 1):
                self._add_question(doc, question, i, question_type, marks_per_q)
        else:
            multi = len(groups) > 1
            for group in groups:
                if multi:
                    self._add_group_subheader(doc, group.get("group_index", 0))
                for i, question in enumerate(group.get("questions", []), 1):
                    self._add_question(
                        doc, question, i, question_type, marks_per_q
                    )

        doc.add_paragraph()

    GROUP_NAMES_AR = [
        "المجموعة الأولى", "المجموعة الثانية", "المجموعة الثالثة",
        "المجموعة الرابعة", "المجموعة الخامسة",
    ]
    GROUP_NAMES_EN = [
        "First group", "Second group", "Third group",
        "Fourth group", "Fifth group",
    ]

    def _add_group_subheader(self, doc: Document, group_index: int) -> None:
        """Render 'First group' / 'Second group' subheader (MOE style)."""
        if self.is_rtl:
            label = (
                self.GROUP_NAMES_AR[group_index]
                if group_index < len(self.GROUP_NAMES_AR)
                else f"المجموعة {group_index + 1}"
            )
        else:
            label = (
                self.GROUP_NAMES_EN[group_index]
                if group_index < len(self.GROUP_NAMES_EN)
                else f"Group {group_index + 1}"
            )
        para = doc.add_paragraph()
        self._set_paragraph_rtl(para)
        para.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT if self.is_rtl else WD_ALIGN_PARAGRAPH.LEFT
        )
        run = para.add_run(f"{label}:")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(80, 80, 80)
        run.font.italic = True

    def _add_question(
        self, doc: Document, question: dict, number: int,
        question_type: str, marks: int,
    ) -> None:
        question_text = question.get(
            "question", question.get("text", f"[Question {number}]")
        )

        q_para = doc.add_paragraph()
        self._set_paragraph_rtl(q_para)
        q_para.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT if self.is_rtl else WD_ALIGN_PARAGRAPH.LEFT
        )

        num_run = q_para.add_run(
            f"({number})  " if self.is_rtl else f"{number}. "
        )
        num_run.font.bold = True
        num_run.font.size = Pt(12)

        text_run = q_para.add_run(question_text)
        text_run.font.size = Pt(12)

        marks_run = q_para.add_run(
            f"    [{marks} درجة]" if self.is_rtl else f"    [{marks}]"
        )
        marks_run.font.size = Pt(9)
        marks_run.font.color.rgb = RGBColor(100, 100, 100)

        if question_type == "multiple_choice":
            options = question.get("options", [])
            self._add_mcq_options(doc, options)
        elif question_type in ("short_answer", "show_work", "long_answer"):
            self._add_answer_space(doc, question_type)

    def _add_mcq_options(self, doc: Document, options: list) -> None:
        if not options:
            return
        letters = (
            ["أ", "ب", "ج", "د"] if self.is_rtl else ["A", "B", "C", "D"]
        )

        if len(options) == 4:
            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for idx, option in enumerate(options[:4]):
                row = idx // 2
                col = idx % 2
                if self.is_rtl:
                    col = 1 - col
                cell = table.cell(row, col)
                cell.width = Cm(8)
                para = cell.paragraphs[0]
                self._set_paragraph_rtl(para)
                run = para.add_run(f"  {letters[idx]}) {option}")
                run.font.size = Pt(11)

            self._remove_table_borders(table)
        else:
            for idx, option in enumerate(options):
                letter = letters[idx] if idx < len(letters) else str(idx + 1)
                opt_para = doc.add_paragraph()
                self._set_paragraph_rtl(opt_para)
                opt_para.paragraph_format.left_indent = Cm(1.5)
                run = opt_para.add_run(f"  {letter}) {option}")
                run.font.size = Pt(11)

    def _add_answer_space(self, doc: Document, question_type: str) -> None:
        lines = {
            "short_answer": 3,
            "show_work": 5,
            "long_answer": 7,
        }.get(question_type, 3)
        for _ in range(lines):
            line_para = doc.add_paragraph()
            line_para.paragraph_format.space_after = Pt(4)
            run = line_para.add_run("." * 90)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(200, 200, 200)

    # ─── Answer Key ──────────────────────────────────────────────────────────

    def _add_answer_key_header(self, doc: Document) -> None:
        title = (
            "نموذج الإجابة ودليل التصحيح"
            if self.is_rtl else "Answer Key & Marking Rubric"
        )
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_rtl(para)
        run = para.add_run(title)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 100, 0)

        subject = self.formatting.get("subject", "الرياضيات")
        grade = self.formatting.get("grade", "")
        info_text = f"{subject} — {grade}" if grade else subject

        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_rtl(info_para)
        info_run = info_para.add_run(info_text)
        info_run.font.size = Pt(12)

        sep_para = doc.add_paragraph()
        sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = sep_para.add_run("═" * 60)
        sep_run.font.size = Pt(10)
        sep_run.font.color.rgb = RGBColor(0, 100, 0)
        doc.add_paragraph()

    def _add_answer_section(self, doc: Document, section: dict) -> None:
        title = section.get("title_ar" if self.is_rtl else "title_en", "")
        marks_per_q = section.get("marks_per_question", 2)

        header_para = doc.add_paragraph()
        self._set_paragraph_rtl(header_para)
        run = header_para.add_run(title)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)

        groups = section.get("groups") or [
            {"group_index": 0, "answers": section.get("answers", [])}
        ]
        multi = len(groups) > 1
        for group in groups:
            answers = group.get("answers", [])
            if not answers:
                continue
            if multi:
                self._add_group_subheader(doc, group.get("group_index", 0))
            self._render_answer_table(doc, answers, marks_per_q)
        doc.add_paragraph()

    def _render_answer_table(
        self, doc: Document, answers: list[dict], marks_per_q: int
    ) -> None:
        table = doc.add_table(rows=len(answers) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = (
            ["رقم", "الإجابة", "خطوات الحل", "الدرجة"]
            if self.is_rtl else ["#", "Answer", "Solution Steps", "Marks"]
        )
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_paragraph_rtl(para)
            run = para.add_run(header)
            run.font.bold = True
            run.font.size = Pt(10)
            self._set_cell_shading(cell, "D5E8D4")

        for idx, answer in enumerate(answers):
            row_idx = idx + 1
            self._set_cell_text(
                table.cell(row_idx, 0),
                str(answer.get("number", idx + 1)), size=10, center=True,
            )
            self._set_cell_text(
                table.cell(row_idx, 1),
                str(answer.get("correct_answer", "")), size=10,
            )
            steps = answer.get("solution_steps", "")
            if isinstance(steps, list):
                steps = "\n".join(steps)
            self._set_cell_text(
                table.cell(row_idx, 2), str(steps)[:200], size=9
            )
            self._set_cell_text(
                table.cell(row_idx, 3),
                str(answer.get("marks", marks_per_q)), size=10, center=True,
            )

        self._style_answer_table(table)

    def _add_marking_rubric(
        self, doc: Document, answer_key: list[dict]
    ) -> None:
        title = (
            "ملخص توزيع الدرجات"
            if self.is_rtl else "Mark Distribution Summary"
        )
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_rtl(para)
        run = para.add_run(title)
        run.font.size = Pt(13)
        run.font.bold = True
        doc.add_paragraph()

        table = doc.add_table(rows=len(answer_key) + 2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = (
            ["السؤال", "عدد الأسئلة", "الدرجة"]
            if self.is_rtl else ["Section", "Questions", "Marks"]
        )
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_paragraph_rtl(para)
            run = para.add_run(header)
            run.font.bold = True
            run.font.size = Pt(10)
            self._set_cell_shading(cell, "D5E8D4")

        total_questions = 0
        total_marks = 0

        for idx, section in enumerate(answer_key):
            row_idx = idx + 1
            # Per-section count uses group 0 only (groups are equivalent
            # so they have the same count) to avoid double-counting.
            primary_answers = section.get("answers")
            if not primary_answers and section.get("groups"):
                primary_answers = section["groups"][0].get("answers", [])
            num_answers = len(primary_answers or [])
            section_marks = num_answers * section.get("marks_per_question", 2)

            title = section.get("title_ar" if self.is_rtl else "title_en", "")
            if ":" in title:
                title = title.split(":")[0].strip()

            self._set_cell_text(table.cell(row_idx, 0), title, size=10)
            self._set_cell_text(
                table.cell(row_idx, 1), str(num_answers), size=10, center=True
            )
            self._set_cell_text(
                table.cell(row_idx, 2), str(section_marks),
                size=10, center=True,
            )
            total_questions += num_answers
            total_marks += section_marks

        total_row = len(answer_key) + 1
        total_label = "المجموع" if self.is_rtl else "Total"
        self._set_cell_text(
            table.cell(total_row, 0), total_label, size=10, bold=True
        )
        self._set_cell_text(
            table.cell(total_row, 1), str(total_questions),
            size=10, center=True, bold=True,
        )
        self._set_cell_text(
            table.cell(total_row, 2), str(total_marks),
            size=10, center=True, bold=True,
        )

        self._style_answer_table(table)

    # ─── Footer ──────────────────────────────────────────────────────────────

    def _add_exam_footer(self, doc: Document) -> None:
        doc.add_paragraph()
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = separator.add_run("═" * 60)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 128)

        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_rtl(footer_para)

        if self.is_rtl:
            footer_text = "مع أطيب التمنيات بالتوفيق والنجاح"
        else:
            footer_text = "Best wishes for success"

        run = footer_para.add_run(footer_text)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0, 0, 128)

    # ─── Helper Methods ──────────────────────────────────────────────────────

    def _set_cell_text(
        self, cell, text: str, bold: bool = False,
        size: int = 11, center: bool = False,
    ) -> None:
        cell.paragraphs[0].clear()
        para = cell.paragraphs[0]
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif self.is_rtl:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._set_paragraph_rtl(para)
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "Arial"

    def _style_header_table(self, table) -> None:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "6")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000080")
            borders.append(border)
        tblPr.append(borders)
        if tbl.tblPr is None:
            tbl.append(tblPr)

    def _style_answer_table(self, table) -> None:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "999999")
            borders.append(border)
        tblPr.append(borders)
        if tbl.tblPr is None:
            tbl.append(tblPr)

    def _remove_table_borders(self, table) -> None:
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "FFFFFF")
            borders.append(border)
        tblPr.append(borders)
        if tbl.tblPr is None:
            tbl.append(tblPr)

    def _set_cell_shading(self, cell, color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), color)
        tcPr.append(shading)
